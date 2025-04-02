# -*- coding: utf-8 -*-
"""
Created on Wed Aug  2 14:04:37 2023

@author: rkf33
"""
import os
import pandas as pd
import secrets
import shutil
import win32com.client
import docx
import nltk


skills = {'ProposedIteration', 'QuantitativeComparison'}

def process_files(dirname,semester,unitSession):
    #save list of filenames
    folders = []
    filenames = []
    for path, subdirs, files in os.walk(dirname):
        for subdir in subdirs:
            folders.append(os.path.join(dirname, subdir))
    for folder in folders:
        for path, subdirs, files in os.walk(folder):
            for name in files:
                filenames.append(os.path.join(folder, name))
        
    #filter out non pdf and non word documents
    extensions = [".pdf", ".docx",".doc"]
    filenames = [f for f in filenames if os.path.splitext(f)[1] in extensions]
    #print(filenames)
    
    if "Consenting" not in dirname:
        non_consent_files = []
        consent = pd.read_excel("Consent.xlsx")
        for i,row in consent.iterrows():
            name = row["First Name"] + " " + row["Last Name"]
            for f in filenames:
                if name in f:
                    non_consent_files.append(f)
        
        filenames = [f for f in filenames if f not in non_consent_files]
    
    n = len(filenames)
    ## extract meta data from filenames
    Semester = pd.Categorical([semester]*n)
    US = pd.Categorical([unitSession]*n)
    
    ID = []
    meta = []
    for i in filenames:
        gen_id = secrets.token_hex(4)
        ID.append(gen_id)
        meta.append(os.path.split(i)[0])
        #copy the pdf with new ID into main folder
        shutil.copy(i,"Processed//" + unitSession + "_" + gen_id + ".pdf")
        #os.rename(os.path.spliti, str(gen_id) + ".pdf")
    #create the dataframe that stores all the metadata
    df = pd.DataFrame(
        { "ID": ID,
          "Filename": meta,
          "Semester": Semester,
          "Unit and Session": US})
    return df


    

def pdf_to_docx(dirname):
    
    word = win32com.client.Dispatch("Word.application")
    
    files = []
    for file in os.listdir(dirname):
        if os.path.splitext(file)[1] == ".pdf":
            files.append(os.path.join(dirname, file))
    for file in files:
        wordDoc = word.Documents.Open(file, False, False, False)
        wordDoc.SaveAs2(os.path.splitext(file)[0])
        wordDoc.Close()
    print(files)

def docx_to_xlsx(dirname):
    for fileloc in os.listdir(dirname):
        if os.path.splitext(fileloc)[1] == ".docx":
            with open(os.path.join(dirname,fileloc),"rb") as file:
                document = docx.Document(file)
                
            raw_text = []
            raw_text_context = []
            raw_tables = []
            
            for table in document.tables:
                for i in range(len(table.rows)):
                    for j in range(len(table.columns)):
                        try:
                            raw_tables.append(table.row_cells(i)[j].text)
                        except:
                            continue
                raw_tables.append("\n")
                 
            for paragraph in document.paragraphs:
                if paragraph.text != "" and paragraph.text != " ":
                    for sentence in nltk.sent_tokenize(paragraph.text):
                        raw_text.append(sentence)
                        raw_text_context.append(paragraph.text)
            file.close()
            with pd.ExcelWriter(os.path.splitext(fileloc)[0] + '.xlsx') as writer:
        
                pd.DataFrame({'Sentences': raw_text} | {skill: [None] * len(raw_text) for skill in skills}).to_excel(
                        writer,
                        sheet_name='Sentence Data')
        
                pd.DataFrame({'Table Cells': raw_tables} | {skill: [None] * len(raw_tables) for skill in skills}).to_excel(
                        writer,
                        sheet_name='Table Data')

#df1 = process_files("U1S1_raw", "F22", "U1S1")
#df2 = process_files("U1S2_raw", "F22", "U1S2")
#df3 = process_files("U2_raw", "F22", "U2")
#df4 = process_files("U3S1_raw", "F22", "U3S1")
#df5 = process_files("U3S2_raw", "F22", "U3S2")
#print("processed F22")

#df6 = process_files("Unit1Session1Notes_ConsentingOnly", "F19", "U1S1")
#df7 = process_files("Unit1Session2Notes_ConsentingOnly", "F19", "U1S2")
#df8 = process_files("Unit2Session1Notes_ConsentingOnly", "F19", "U2S1")
#df9 = process_files("Unit2Session2Notes_ConsentingOnly", "F19", "U2S2")
#df10 = process_files("Unit3Session1Notes_ConsentingOnly", "F19", "U3S1")
#df11 = process_files("Unit3Session2Notes_ConsentingOnly", "F19", "U3S2")
#print("processed F19")


#df = pd.concat([df1,df2,df3,df4,df5,df6,df7,df8,df9,df10,df11])
#df.to_csv("Metadata.csv")

#pdf_to_docx("C:\\Users\\rkf33\\Documents\\Lab_notes_processing\\Processed")
#docx_to_xlsx("C:\\Users\\rkf33\\Documents\\Lab_notes_processing\\Processed")


